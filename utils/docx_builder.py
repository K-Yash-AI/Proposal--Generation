"""
DOCX builder for proposal generation.

Two modes:
  1. Template mode  – opens an existing .docx (exported from Google Doc),
                       finds placeholder tokens, and replaces them with AI-
                       generated content.
  2. Default mode   – builds a brand-new, professionally formatted .docx
                       from scratch when no template is provided.
"""
from __future__ import annotations

import io
import re
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

from utils.logger import log

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
_BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)   # near-black navy
_BRAND_ACCENT = RGBColor(0xE9, 0x4F, 0x37)  # Frutitions orange-red
_GREY = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_run_color(run, color: RGBColor) -> None:
    run.font.color.rgb = color


def _add_horizontal_rule(doc: Document) -> None:
    """Insert a thin horizontal line (using paragraph border)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _add_section_heading(doc: Document, title: str) -> None:
    h = doc.add_heading(title, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.color.rgb = _BRAND_DARK
        run.font.size = Pt(16)
    _add_horizontal_rule(doc)


def _add_body_text(doc: Document, text: str) -> None:
    """Add multi-paragraph body text, preserving blank-line breaks."""
    paragraphs = text.strip().split("\n\n")
    for para_text in paragraphs:
        # Handle single-newline line breaks within a paragraph
        lines = para_text.strip().split("\n")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        first = True
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if not first:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.color.rgb = _GREY
            first = False


def _add_table_from_text(doc: Document, text: str) -> None:
    """
    Parse a markdown-style table from text and render it as a Word table.
    Falls back to body text if parsing fails.
    """
    lines = [l.strip() for l in text.strip().splitlines() if l.strip()]
    table_lines = [l for l in lines if l.startswith("|")]

    if len(table_lines) < 2:
        _add_body_text(doc, text)
        return

    try:
        def parse_row(line: str) -> list[str]:
            return [c.strip() for c in line.strip("|").split("|")]

        headers = parse_row(table_lines[0])
        # Skip separator row (---|---|---)
        data_rows = [parse_row(l) for l in table_lines[2:] if "---" not in l]

        tbl = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
        tbl.style = "Table Grid"

        # Header row
        hdr_row = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Dark background
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1A1A2E")
            tcPr.append(shd)

        for r_idx, row_data in enumerate(data_rows, start=1):
            row = tbl.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < len(headers):
                    row.cells[c_idx].text = cell_text

        doc.add_paragraph()  # spacing after table

    except Exception:
        _add_body_text(doc, text)


# ---------------------------------------------------------------------------
# Core builder
# ---------------------------------------------------------------------------


class DocxBuilder:
    """Builds the final proposal DOCX."""

    def __init__(self, template_bytes: Optional[bytes] = None) -> None:
        self._template_bytes = template_bytes

    # ── Template-replacement mode ────────────────────────────────────────────

    def build_from_template(
        self,
        sections: dict[str, str],
        metadata: dict[str, str],
        output_path: str,
    ) -> str:
        """
        Open the template DOCX, replace all placeholder tokens, and save.

        Args:
            sections:   mapping of section_key → generated content
            metadata:   mapping of meta_key → value (client_name, date, etc.)
            output_path: where to write the output .docx

        Returns:
            Absolute path of the written file.
        """
        assert self._template_bytes, "Template bytes required for template mode"

        doc = Document(io.BytesIO(self._template_bytes))
        replacements = self._build_replacements(sections, metadata)

        log.info("[step]Injecting content into template placeholders[/step]")
        self._replace_in_document(doc, replacements)

        out = Path(output_path)
        doc.save(str(out))
        log.info(f"[success]Saved proposal → {out.resolve()}[/success]")
        return str(out.resolve())

    @staticmethod
    def _build_replacements(
        sections: dict[str, str], metadata: dict[str, str]
    ) -> dict[str, str]:
        replacements: dict[str, str] = {}
        for key, value in sections.items():
            token = f"{{{{{key.upper()}}}}}"
            replacements[token] = value
        for key, value in metadata.items():
            token = f"{{{{{key.upper()}}}}}"
            replacements[token] = value
        return replacements

    @staticmethod
    def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
        """Replace tokens inside a paragraph, merging runs if necessary."""
        full_text = "".join(r.text for r in paragraph.runs)
        modified = full_text
        for token, value in replacements.items():
            modified = modified.replace(token, value)

        if modified != full_text:
            # Clear all runs and put content in first run
            for i, run in enumerate(paragraph.runs):
                run.text = modified if i == 0 else ""

    def _replace_in_document(self, doc: Document, replacements: dict[str, str]) -> None:
        # Paragraphs in body
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        # Paragraphs inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        # Headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for para in header.paragraphs:
                        self._replace_in_paragraph(para, replacements)
            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        self._replace_in_paragraph(para, replacements)

    # ── Default (from-scratch) mode ──────────────────────────────────────────

    def build_default(
        self,
        sections: dict[str, str],
        metadata: dict[str, str],
        output_path: str,
    ) -> str:
        """
        Build a fully formatted proposal DOCX from scratch.

        Args:
            sections:   mapping of section_key → generated content
            metadata:   client / proposal metadata
            output_path: where to write

        Returns:
            Absolute path of the written file.
        """
        log.info("[step]Building proposal DOCX from default template[/step]")
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        self._build_cover_page(doc, metadata)
        doc.add_page_break()
        self._build_toc_placeholder(doc)
        doc.add_page_break()
        self._build_sections(doc, sections)
        self._add_footer(doc, metadata)

        out = Path(output_path)
        doc.save(str(out))
        log.info(f"[success]Saved proposal → {out.resolve()}[/success]")
        return str(out.resolve())

    def _build_cover_page(self, doc: Document, metadata: dict[str, str]) -> None:
        from config import settings

        # Company name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(80)
        run = p.add_run(settings.company_name.upper())
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = _BRAND_ACCENT

        # Tagline
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(settings.company_tagline)
        run.font.size = Pt(14)
        run.font.color.rgb = _GREY
        run.font.italic = True

        doc.add_paragraph()
        _add_horizontal_rule(doc)
        doc.add_paragraph()

        # Proposal title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PROPOSAL")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = _BRAND_DARK

        # Client name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client = metadata.get("client_company") or metadata.get("client_name", "")
        run = p.add_run(f"Prepared for: {client}")
        run.font.size = Pt(16)
        run.font.color.rgb = _BRAND_DARK

        doc.add_paragraph()

        # Date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(metadata.get("proposal_date", str(date.today().strftime("%B %d, %Y"))))
        run.font.size = Pt(12)
        run.font.color.rgb = _GREY

        # Prepared by
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Prepared by: {metadata.get('prepared_by', settings.company_name)}")
        run.font.size = Pt(12)
        run.font.color.rgb = _GREY

        doc.add_paragraph()
        _add_horizontal_rule(doc)

        # Contact info
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"{settings.company_website}  |  {settings.company_email}  |  {settings.company_phone}"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = _GREY

    def _build_toc_placeholder(self, doc: Document) -> None:
        h = doc.add_heading("Table of Contents", level=1)
        for run in h.runs:
            run.font.color.rgb = _BRAND_DARK
        p = doc.add_paragraph()
        run = p.add_run(
            "[This table of contents is auto-generated by Word. "
            'Right-click and select "Update Field" to refresh after editing.]'
        )
        run.font.italic = True
        run.font.color.rgb = _GREY
        run.font.size = Pt(10)

    def _build_sections(self, doc: Document, sections: dict[str, str]) -> None:
        from templates.proposal_sections import PROPOSAL_SECTIONS

        for section_def in PROPOSAL_SECTIONS:
            content = sections.get(section_def.key, "")
            if not content:
                continue

            _add_section_heading(doc, section_def.title)

            # Timeline section: try rendering as table
            if section_def.key == "project_timeline" and "|" in content:
                _add_table_from_text(doc, content)
            else:
                _add_body_text(doc, content)

    def _add_footer(self, doc: Document, metadata: dict[str, str]) -> None:
        from config import settings

        for sec in doc.sections:
            footer = sec.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(
                f"Confidential | {settings.company_name} | "
                f"{settings.company_website} | "
                f"© {date.today().year}"
            )
            run.font.size = Pt(8)
            run.font.color.rgb = _GREY
