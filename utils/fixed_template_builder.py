"""
Fixed Template DOCX Builder — Fruition Enterprise Proposal Format.

Builds a professional, branded DOCX matching the Fruition PDF template:
  Page 1  — Cover page
  Page 2  — Overview boxes (01 / 02 / 03)
  Page 3  — Landscape & Objective Overview
  Page 4  — Solution Design + embedded workflow diagram
  Page 5+ — Phase detail pages (1, 2, 3)
  Next    — Project Deliverables & Timeline table
  Next    — Investment Summary tables
  Next    — Next Steps, Team Contacts, Future Opportunities
  Final   — Fruition Introduction (boilerplate)

All content (except boilerplate) is dynamically injected from AI-generated
sections and Python-computed pricing — zero hardcoded dummy data.
"""
from __future__ import annotations

import io
import json
import re
from datetime import date
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from utils.diagram_generator import generate_workflow_diagram, generate_fallback_diagram
from utils.logger import log

# ── Colour constants ──────────────────────────────────────────────────────────
_PURPLE = RGBColor(0x5B, 0x2D, 0x8F)       # #5B2D8F  primary purple
_PURPLE_MID = RGBColor(0x7B, 0x4D, 0xAF)   # #7B4DAF
_PURPLE_LIGHT = RGBColor(0xE8, 0xE0, 0xF0) # #E8E0F0  light purple fill
_DARK = RGBColor(0x2D, 0x2D, 0x2D)         # near-black body text
_GREY = RGBColor(0x7F, 0x7F, 0x7F)         # secondary text
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Hex strings for XML shading
_PURPLE_HEX = "5B2D8F"
_PURPLE_LIGHT_HEX = "E8E0F0"
_LIGHT_GREY_HEX = "F5F5F5"
_WHITE_HEX = "FFFFFF"


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_border(cell, side: str = "bottom", color: str = "E0E0E0", sz: str = "4") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), sz)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)
    tcBorders.append(border)
    tcPr.append(tcBorders)


def _fmt(run, size: float, color: RGBColor, bold: bool = False, italic: bool = False) -> None:
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def _heading(doc: Document, text: str, level: int, color: RGBColor = _DARK,
             size: float = 20, space_before: float = 12, space_after: float = 8) -> None:
    """Add a heading paragraph without any horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    sizes = {1: 22, 2: 15, 3: 12}
    run = p.add_run(text)
    _fmt(run, sizes.get(level, size), color, bold=True)


def _body(doc: Document, text: str, size: float = 11, color: RGBColor = _DARK,
          space_after: float = 8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _fmt(run, size, color)


def _spacer(doc: Document, pts: float = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)


# ── Inline bold markdown renderer ─────────────────────────────────────────────

def _add_inline(paragraph, text: str, base_size: float = 11,
                base_color: RGBColor = _DARK) -> None:
    """Add text to paragraph with **bold** markdown support."""
    # Clean common AI artefacts
    text = text.replace("**:", ":**").replace("––", "—")
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        _fmt(run, base_size, base_color, bold=(idx % 2 == 1))


# ── Section content renderer ──────────────────────────────────────────────────

def _render_content(doc: Document, content: str) -> None:
    """
    General-purpose renderer for AI-generated section content.

    Handles:
      • Lines starting with '•', '-', '*' → bullet points with inline bold
      • Lines matching '\\d+\\. text'       → numbered sub-headings (purple, bold)
      • Short lines ending in ':'          → sub-headings (dark, bold)
      • Blank lines                        → paragraph separator
      • Everything else                    → body paragraph with inline bold
    """
    # Strip section markers the AI may have included literally
    _STRIP_MARKERS = [
        "COMPANY_OVERVIEW:", "CONTEXT:", "CURRENT_CHALLENGES:", "GOAL:",
        "OVERVIEW_TEXT:", "ARCHITECTURE_DESCRIPTION:", "DIAGRAM_JSON:",
        "PHASE_NAME:", "INTRO:",
    ]

    lines = content.split("\n")
    para_buf: list[str] = []

    def flush_para():
        if para_buf:
            joined = " ".join(l for l in para_buf if l)
            if joined:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                _add_inline(p, joined)
            para_buf.clear()

    for raw_line in lines:
        line = raw_line.strip()

        # Skip section markers
        if any(line.upper() == m.upper() or line.upper().startswith(m.upper()) and len(line) < len(m) + 3
               for m in _STRIP_MARKERS):
            flush_para()
            continue

        # Blank line → flush buffer
        if not line:
            flush_para()
            continue

        # Numbered sub-heading: "1. Short Title" (not a full sentence)
        num_match = re.match(r"^(\d+)\.\s+(.+)", line)
        if num_match:
            text = num_match.group(2)
            # Treat as sub-heading if short & doesn't end like a sentence
            if len(text) <= 70 and not text[-1] in ".!?" or text.endswith(":"):
                flush_para()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(f"{num_match.group(1)}. {text.rstrip(':')}")
                _fmt(run, 11.5, _PURPLE, bold=True)
                continue
            # Long numbered item → treat as bullet
            flush_para()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            _add_inline(p, f"{num_match.group(1)}. {text}")
            continue

        # Bullet item: •, -, *
        if re.match(r"^[•\-\*–]\s*", line):
            flush_para()
            text = re.sub(r"^[•\-\*–]\s*", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(4)
            # Add bullet char as plain text, then the rest with inline bold
            run = p.add_run("• ")
            _fmt(run, 11, _DARK)
            _add_inline(p, text)
            continue

        # Sub-heading: short line ending in ":"
        if line.endswith(":") and 3 < len(line) <= 80 and not line.startswith("•"):
            flush_para()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line.rstrip(":"))
            _fmt(run, 11.5, _DARK, bold=True)
            continue

        # Regular text → accumulate
        para_buf.append(line)

    flush_para()


# ── Section parsers ───────────────────────────────────────────────────────────

def _parse_landscape(content: str) -> dict[str, str]:
    """
    Split landscape_objective into its 4 parts using marker detection.
    Returns dict with keys: company_overview, context, challenges, goal.
    """
    markers = {
        "COMPANY_OVERVIEW:": "company_overview",
        "CONTEXT:": "context",
        "CURRENT_CHALLENGES:": "challenges",
        "GOAL:": "goal",
    }
    result: dict[str, list[str]] = {v: [] for v in markers.values()}
    current = "company_overview"

    for line in content.split("\n"):
        matched = False
        for marker, key in markers.items():
            if line.strip().upper().startswith(marker.upper()):
                current = key
                # Remainder of the line after the marker
                rest = line.strip()[len(marker):].strip()
                if rest:
                    result[current].append(rest)
                matched = True
                break
        if not matched:
            result[current].append(line)

    return {k: "\n".join(v).strip() for k, v in result.items()}


def _parse_phase_name(content: str, default: str) -> tuple[str, str]:
    """Extract PHASE_NAME from content. Returns (name, remaining_content)."""
    lines = content.split("\n")
    name = default
    rest_start = 0
    for i, line in enumerate(lines):
        if line.strip().upper().startswith("PHASE_NAME:"):
            candidate = line.strip()[len("PHASE_NAME:"):].strip()
            if candidate:
                name = candidate
                rest_start = i + 1
            break
    return name, "\n".join(lines[rest_start:]).strip()


def _parse_solution_design(content: str) -> tuple[str, Optional[dict]]:
    """
    Split solution_design into text and diagram spec.
    Returns (text_content, diagram_dict_or_None).
    """
    diagram_json = None
    text_part = content

    if "DIAGRAM_JSON:" in content:
        parts = content.split("DIAGRAM_JSON:", 1)
        text_part = parts[0].strip()
        json_raw = parts[1].strip()
        json_raw = re.sub(r"```(?:json)?", "", json_raw).strip().strip("`").strip()
        try:
            diagram_json = json.loads(json_raw)
        except json.JSONDecodeError:
            # Try to extract just the JSON object
            m = re.search(r"\{[\s\S]+\}", json_raw)
            if m:
                try:
                    diagram_json = json.loads(m.group())
                except Exception:
                    pass

    # Strip OVERVIEW_TEXT: marker from text
    text_part = re.sub(r"(?i)^OVERVIEW_TEXT:\s*", "", text_part, flags=re.MULTILINE).strip()
    text_part = re.sub(r"(?i)^ARCHITECTURE_DESCRIPTION:\s*", "", text_part, flags=re.MULTILINE).strip()
    return text_part, diagram_json


# ── Table helpers ─────────────────────────────────────────────────────────────

def _purple_header_row(table, headers: list[str], col_widths: Optional[list[float]] = None) -> None:
    """Style the first row of a table as a purple header."""
    row = table.rows[0]
    for i, cell in enumerate(row.cells):
        cell.text = ""
        _set_cell_bg(cell, _PURPLE_HEX)
        if col_widths and i < len(col_widths):
            cell.width = Inches(col_widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(headers[i])
        _fmt(run, 10, _WHITE, bold=True)


def _add_table_row(table, values: list[str], bg_hex: str = _WHITE_HEX,
                   bold: bool = False, color: RGBColor = _DARK) -> None:
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        _set_cell_bg(cell, bg_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        _fmt(run, 10, color, bold=bold)


# ── Main builder ──────────────────────────────────────────────────────────────

class FixedTemplateBuilder:
    """Builds the complete Fruition enterprise proposal DOCX."""

    def build(
        self,
        sections: dict[str, str],
        metadata: dict[str, Any],
        output_path: str,
    ) -> str:
        log.info("[step]Building proposal document (Fruition enterprise format)[/step]")
        doc = Document()

        # Page margins
        for sec in doc.sections:
            sec.top_margin = Inches(0.8)
            sec.bottom_margin = Inches(0.8)
            sec.left_margin = Inches(1.0)
            sec.right_margin = Inches(1.0)

        # ── Pages ─────────────────────────────────────────────────────────────
        self._cover_page(doc, metadata)
        doc.add_page_break()

        self._overview_boxes(doc)
        doc.add_page_break()

        self._landscape_objective_page(doc, sections, metadata)
        doc.add_page_break()

        self._solution_design_page(doc, sections, metadata)
        doc.add_page_break()

        self._phase_page(
            doc,
            sections.get("phase_1_content", ""),
            metadata["pricing"],
            phase_num=1,
        )
        doc.add_page_break()

        self._phase_page(
            doc,
            sections.get("phase_2_content", ""),
            metadata["pricing"],
            phase_num=2,
        )
        doc.add_page_break()

        self._phase_page(
            doc,
            sections.get("phase_3_content", ""),
            metadata["pricing"],
            phase_num=3,
        )
        doc.add_page_break()

        self._deliverables_timeline_page(doc, metadata)
        doc.add_page_break()

        self._investment_summary_page(doc, sections, metadata)
        doc.add_page_break()

        self._next_steps_page(doc, sections, metadata)
        doc.add_page_break()

        self._fruition_introduction(doc)

        # Footer
        self._add_footer(doc, metadata)

        out = Path(output_path)
        doc.save(str(out))
        log.info(f"[success]Saved → {out.resolve()}[/success]")
        return str(out.resolve())

    # ── Page 1: Cover ─────────────────────────────────────────────────────────

    def _cover_page(self, doc: Document, metadata: dict[str, Any]) -> None:
        for _ in range(4):
            _spacer(doc, 8)

        # "Proposal" — large purple title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run("Proposal")
        _fmt(run, 72, _PURPLE, bold=True)

        # Client company name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        company = metadata.get("client_company") or metadata.get("client_name", "")
        run = p.add_run(company)
        _fmt(run, 24, _DARK, bold=True)

        _spacer(doc, 12)

        # "Presented by:"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("Presented by: ")
        _fmt(run, 12, _GREY)
        run2 = p.add_run(metadata.get("prepared_by", metadata.get("company_name", "Fruition")))
        _fmt(run2, 12, _DARK, bold=True)

        # Date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(metadata.get("proposal_date", date.today().strftime("%B %dth, %Y")))
        _fmt(run, 12, _GREY)

    # ── Page 2: Overview boxes (01 / 02 / 03) ─────────────────────────────────

    def _overview_boxes(self, doc: Document) -> None:
        _spacer(doc, 8)
        self._overview_box(
            doc, "01",
            "Landscape & Objective",
            "We assess your current state, identify pain points, and define clear objectives "
            "for your monday.com implementation. This discovery phase ensures we understand your "
            "unique challenges and align on success metrics before we begin.",
        )
        _spacer(doc, 16)
        self._overview_box(
            doc, "02",
            "Solution Design",
            "Our tailored solution blueprint addresses your specific workflows through board "
            "architecture, automation strategies, and integrations. This section details the "
            "complete scope of work, phase breakdowns, and deliverables for your implementation.",
        )
        _spacer(doc, 16)
        self._overview_box(
            doc, "03",
            "Summary",
            "A transparent breakdown of investment, timelines, and resources required to bring "
            "your solution to life. Here you'll find the total engagement hours, costs, and "
            "what's included in each implementation phase.",
        )

    def _overview_box(self, doc: Document, number: str, title: str, description: str) -> None:
        # Large purple number
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(number)
        _fmt(run, 36, _PURPLE, bold=True)

        # Title
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        _fmt(run, 13, _DARK, bold=True)

        # Description
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(description)
        _fmt(run, 11, _GREY)

    # ── Page 3: Landscape & Objective ─────────────────────────────────────────

    def _landscape_objective_page(
        self, doc: Document, sections: dict[str, str], metadata: dict[str, Any]
    ) -> None:
        _heading(doc, "Landscape & Objective Overview", level=1, color=_DARK,
                 space_before=0, space_after=14)

        content = sections.get("landscape_objective", "")
        if not content:
            _body(doc, "Discovery insights will be documented here following initial workshops.")
            return

        parsed = _parse_landscape(content)
        client_company = metadata.get("client_company", "")

        # Company overview — bold the company name in the first sentence
        overview = parsed.get("company_overview", "")
        if overview:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            # Bold company name at start if found
            if overview.startswith(client_company):
                run_bold = p.add_run(client_company)
                _fmt(run_bold, 11, _DARK, bold=True)
                rest = overview[len(client_company):]
                _add_inline(p, rest)
            else:
                _add_inline(p, overview)

        # Context paragraph
        context = parsed.get("context", "")
        if context:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            _add_inline(p, context)

        # Current Challenges heading
        challenges = parsed.get("challenges", "")
        if challenges:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("Current Challenges")
            _fmt(run, 12, _PURPLE, bold=True)

            for line in challenges.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # Strip leading bullet characters
                text = re.sub(r"^[•\-\*]\s*", "", line)
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Inches(0.25)
                bp.paragraph_format.first_line_indent = Inches(-0.18)
                bp.paragraph_format.space_after = Pt(5)
                run_dot = bp.add_run("• ")
                _fmt(run_dot, 11, _DARK)
                _add_inline(bp, text)

        # Goal paragraph
        goal = parsed.get("goal", "")
        if goal:
            _spacer(doc, 8)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            _add_inline(p, goal)

    # ── Page 4: Solution Design ────────────────────────────────────────────────

    def _solution_design_page(
        self, doc: Document, sections: dict[str, str], metadata: dict[str, Any]
    ) -> None:
        _heading(doc, "Solution Design", level=1, color=_DARK, space_before=0, space_after=12)

        content = sections.get("solution_design", "")
        text_part, diagram_spec = _parse_solution_design(content)

        # Render overview text first
        if text_part:
            _render_content(doc, text_part)

        # Generate and embed workflow diagram
        _spacer(doc, 10)
        try:
            if diagram_spec:
                log.info("[step]Generating workflow diagram...[/step]")
                png_bytes = generate_workflow_diagram(diagram_spec)
                caption = diagram_spec.get("caption", "")
            else:
                log.info("[step]No diagram spec — generating fallback diagram[/step]")
                # Build fallback from pipeline stages in insights or generic
                fallback_stages = ["Discovery", "Configuration", "Automation", "Testing", "Go-Live"]
                png_bytes = generate_fallback_diagram(fallback_stages)
                caption = "End-to-end implementation lifecycle"

            img_stream = io.BytesIO(png_bytes)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            run.add_picture(img_stream, width=Inches(6.0))

        except Exception as exc:
            log.warning(f"Diagram generation failed: {exc} — skipping diagram.")
            caption = ""

        # Caption
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(caption)
            _fmt(run, 9, _GREY, italic=True)

    # ── Phase pages ───────────────────────────────────────────────────────────

    def _phase_page(
        self,
        doc: Document,
        content: str,
        pricing: dict[str, Any],
        phase_num: int,
    ) -> None:
        key = f"phase_{phase_num}"
        default_name = pricing.get(f"phase_{phase_num}_name", f"Phase {phase_num}")
        hours = pricing.get(f"phase_{phase_num}_hours", 10)
        weeks = pricing.get(f"phase_{phase_num}_weeks", "TBD")

        phase_name, body_content = _parse_phase_name(content, default_name)

        # Phase heading: "Phase 1: Discovery & Requirements Workshops"
        _heading(
            doc,
            f"Phase {phase_num}: {phase_name}",
            level=1,
            color=_DARK,
            space_before=0,
            space_after=6,
        )

        # Hours + Timeline line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(14)
        run_hours = p.add_run(f"Hours: {hours}    ")
        _fmt(run_hours, 11, _GREY, bold=True)
        run_tl = p.add_run(f"Timeline: Weeks {weeks}")
        _fmt(run_tl, 11, _GREY, bold=True)

        if body_content:
            _render_content(doc, body_content)
        else:
            _body(doc, "Phase details will be documented following discovery workshops.")

    # ── Project Deliverables & Timeline ───────────────────────────────────────

    def _deliverables_timeline_page(
        self, doc: Document, metadata: dict[str, Any]
    ) -> None:
        pricing = metadata["pricing"]

        _heading(doc, "Project Deliverables & Timeline", level=1, color=_DARK,
                 space_before=0, space_after=12)

        # Summary line
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        run = p.add_run(
            f"Project Duration: {pricing['total_weeks']} Weeks    "
            f"Total Hours: {pricing['total_hours']} Hours    "
            f"Initial Users: ~{pricing['initial_users']}"
        )
        _fmt(run, 11, _GREY, bold=True)

        # Table: Phase | Week | Deliverables | Hours | Responsible
        headers = ["Phase", "Week", "Deliverables", "Hours", "Responsible"]
        col_widths = [1.3, 0.9, 2.6, 0.7, 1.1]
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        _purple_header_row(table, headers, col_widths)

        rows = [
            (
                f"Phase 1: {pricing['phase_1_name']}",
                f"Weeks {pricing['phase_1_weeks']}",
                pricing["phase_1_deliverables"],
                str(pricing["phase_1_hours"]),
                pricing["phase_1_lead"],
            ),
            (
                f"Phase 2: {pricing['phase_2_name']}",
                f"Weeks {pricing['phase_2_weeks']}",
                pricing["phase_2_deliverables"],
                str(pricing["phase_2_hours"]),
                pricing["phase_2_lead"],
            ),
            (
                f"Phase 3: {pricing['phase_3_name']}",
                f"Weeks {pricing['phase_3_weeks']}",
                pricing["phase_3_deliverables"],
                str(pricing["phase_3_hours"]),
                pricing["phase_3_lead"],
            ),
        ]

        for i, row_data in enumerate(rows):
            bg = _LIGHT_GREY_HEX if i % 2 == 0 else _WHITE_HEX
            _add_table_row(table, list(row_data), bg_hex=bg)

        # Total row
        _add_table_row(
            table,
            ["Total", f"{pricing['total_weeks']} Weeks", "", str(pricing["total_hours"]), ""],
            bg_hex=_PURPLE_LIGHT_HEX,
            bold=True,
            color=_PURPLE,
        )

        _spacer(doc, 8)

    # ── Investment Summary ─────────────────────────────────────────────────────

    def _investment_summary_page(
        self, doc: Document, sections: dict[str, str], metadata: dict[str, Any]
    ) -> None:
        pricing = metadata["pricing"]
        currency = pricing.get("currency", "AUD")

        _heading(doc, "Investment Summary", level=1, color=_DARK, space_before=0, space_after=14)

        # ── Year 1 ────────────────────────────────────────────────────────────
        _heading(doc, f"Year 1 Investment ({currency})", level=2, color=_DARK,
                 space_before=6, space_after=8)

        headers_inv = ["Item", f"Cost ({currency})", "Notes"]
        col_w_inv = [2.8, 1.2, 2.6]
        tbl1 = doc.add_table(rows=1, cols=3)
        tbl1.style = "Table Grid"
        _purple_header_row(tbl1, headers_inv, col_w_inv)

        # monday.com licensing row
        products_str = " + ".join(pricing.get("monday_products", ["Work Management"]))
        users = pricing["initial_users"]
        lic_desc = pricing.get("licensing_year1_description") or \
                   f"{users} users — {products_str}"
        _add_table_row(tbl1, [
            f"monday.com Licensing ({users} users)",
            pricing["licensing_year1_fmt"],
            lic_desc,
        ])

        # Implementation phases
        for ph in (1, 2, 3):
            _add_table_row(tbl1, [
                f"Implementation — Phase {ph}: {pricing[f'phase_{ph}_name']}",
                pricing[f"phase_{ph}_cost_fmt"],
                f"{pricing[f'phase_{ph}_hours']} hours @ ${pricing['hourly_rate']}/hr",
            ], bg_hex=_LIGHT_GREY_HEX)

        # Total row
        _add_table_row(
            tbl1,
            ["Total Year 1 Investment", pricing["total_year1_fmt"], "Licensing + Implementation"],
            bg_hex=_PURPLE_LIGHT_HEX,
            bold=True,
            color=_PURPLE,
        )

        _spacer(doc, 12)

        # ── Year 2+ ───────────────────────────────────────────────────────────
        _heading(doc, f"Year 2+ Ongoing ({currency})", level=2, color=_DARK,
                 space_before=6, space_after=8)

        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.style = "Table Grid"
        _purple_header_row(tbl2, headers_inv, col_w_inv)

        lic2_desc = pricing.get("licensing_year2_description") or \
                    f"Annual platform licensing — {users} users"
        _add_table_row(tbl2, [
            "monday.com Licensing (ongoing)",
            pricing["licensing_year2_fmt"],
            lic2_desc,
        ])
        _add_table_row(tbl2, [
            "Managed Services (optional)",
            pricing["managed_services_monthly_fmt"] + " / month",
            "Ongoing support, optimisation, new workflows as needed",
        ], bg_hex=_LIGHT_GREY_HEX)

        _spacer(doc, 14)

        # Notes section
        notes_content = sections.get("investment_notes", "")
        if notes_content:
            _heading(doc, "Notes", level=2, color=_DARK, space_before=4, space_after=6)
            _render_content(doc, notes_content)

    # ── Next Steps, Team Contacts, Future Opportunities ───────────────────────

    def _next_steps_page(
        self, doc: Document, sections: dict[str, str], metadata: dict[str, Any]
    ) -> None:
        pricing = metadata["pricing"]

        # ── Next Steps ────────────────────────────────────────────────────────
        _heading(doc, "Next Steps", level=1, color=_DARK, space_before=0, space_after=12)

        next_content = sections.get("next_steps", "")
        if next_content:
            _render_content(doc, next_content)
        else:
            for i, step in enumerate([
                "Review this proposal and confirm the phased approach aligns with your requirements and budget",
                "Confirm licensing direction and schedule a call with the monday.com account team",
                "Sign off on Phase 1 scope and schedule the first Discovery & Requirements Workshop",
                "Prepare existing process documentation (workflows, data exports, system access) for Workshop 1",
                f"{settings.company_name if hasattr(settings, 'company_name') else 'Fruition'} to set up workspace and commence build following Phase 1 sign-off",
            ], 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(f"{i}. {step}")
                _fmt(run, 11, _DARK)

        _spacer(doc, 16)

        # ── Project Team Contacts ─────────────────────────────────────────────
        _heading(doc, "Project Team Contacts", level=2, color=_DARK, space_before=8, space_after=8)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _purple_header_row(tbl, [
            f"{metadata.get('company_name', 'Fruition')} Services",
            f"{metadata.get('client_company', 'Client')}",
        ])

        # Fruition contacts
        fruition_text = (
            f"{metadata.get('prepared_by', 'Project Lead')}\n"
            f"{metadata.get('company_email', '')}\n"
            f"{metadata.get('company_phone', '')}"
        )

        # Client contacts from key_contacts
        key_contacts: list[dict] = metadata.get("key_contacts", [])
        if key_contacts:
            client_text = "\n".join(
                f"{c.get('name', '')} — {c.get('role', '')}"
                for c in key_contacts[:3]
            )
        else:
            client_text = (
                f"{metadata.get('client_name', 'Client Contact')}\n"
                f"{metadata.get('client_email', '')}"
            )

        row = tbl.add_row()
        for i, text in enumerate([fruition_text, client_text]):
            cell = row.cells[i]
            _set_cell_bg(cell, _LIGHT_GREY_HEX)
            for j, line in enumerate(text.split("\n")):
                if j == 0:
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(line)
                _fmt(run, 10, _DARK, bold=(j == 0))

        _spacer(doc, 20)

        # ── Future Phase Opportunities ────────────────────────────────────────
        _heading(doc, "Future Phase Opportunities", level=2, color=_DARK,
                 space_before=8, space_after=8)

        future_content = sections.get("future_opportunities", "")
        if future_content:
            _render_content(doc, future_content)
        else:
            defaults = [
                "Platform API integration with existing systems",
                "AI-powered automation and predictive analytics",
                "Advanced reporting and executive dashboards",
                "Mobile app enhancement and offline capability",
                "Multi-entity or multi-site expansion rollout",
            ]
            for i, item in enumerate(defaults, 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(5)
                run = p.add_run(f"{i}. {item}")
                _fmt(run, 11, _DARK)

    # ── Final page: Fruition Introduction (KEEP EXACTLY AS IS) ───────────────

    def _fruition_introduction(self, doc: Document) -> None:
        _heading(doc, "FRUITION INTRODUCTION", level=1, color=_DARK, space_before=0, space_after=14)

        # Bold tagline
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        run = p.add_run(
            "monday.com Advanced Implementation and Integration Partner. "
            "Trusted by 600+ Private and Public organisations globally."
        )
        _fmt(run, 12, _DARK, bold=True)

        boilerplate = [
            (
                "As a certified monday.com partner, we bring deep expertise to every implementation. "
                "We begin by listening carefully to understand your unique needs and objectives, then "
                "collaborate with you to design solutions that fit seamlessly into your workflows."
            ),
            (
                "Our experienced team guides you through each phase, maintaining clear communication "
                "and delivering on time and within budget. Beyond implementation, we provide ongoing "
                "support to ensure your continued success."
            ),
            (
                "At Fruition, we measure our success by yours — building lasting partnerships that "
                "help your business thrive."
            ),
        ]
        for para_text in boilerplate:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(para_text)
            _fmt(run, 11, _DARK)

    # ── Footer ────────────────────────────────────────────────────────────────

    def _add_footer(self, doc: Document, metadata: dict[str, Any]) -> None:
        from config import settings as cfg
        for sec in doc.sections:
            footer = sec.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(
                f"Confidential  |  {cfg.company_name}  |  "
                f"{cfg.company_website}  |  © {date.today().year}"
            )
            _fmt(run, 8, _GREY)
